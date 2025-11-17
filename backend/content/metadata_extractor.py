"""
AI-powered metadata extraction service using Gemini
Optimized for performance, memory efficiency, and reliability
"""
import json
import os
import re
import tempfile
import time

from django.conf import settings
from google import genai
from google.genai import types


class MetadataExtractor:
    """Service for extracting metadata from files using Gemini AI"""

    # Constants for optimization
    MAX_POLLING_ATTEMPTS = 60  # Maximum 60 seconds wait
    INITIAL_POLL_INTERVAL = 0.5  # Start with 500ms
    MAX_POLL_INTERVAL = 5.0  # Cap at 5 seconds
    POLL_BACKOFF_MULTIPLIER = 1.5  # Exponential backoff
    JSON_PATTERN = re.compile(r'```(?:json)?\s*(.*?)\s*```', re.DOTALL)

    def __init__(self):
        if not settings.GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY not configured")
        self.client = genai.Client(api_key=settings.GEMINI_API_KEY)
        self.model = settings.GEMINI_MODEL

    def extract_metadata(self, file) -> dict:
        """
        Extract metadata from a file using Gemini AI.
        Returns a dictionary with title, description, subject, difficulty, author, publication_year
        Handles DOCX by converting to text first (Gemini doesn't support DOCX directly)
        """
        tmp_file_path = None
        uploaded_file = None
        is_docx = False
        
        try:
            # Check if file is DOCX (Gemini doesn't support DOCX directly)
            file_name_lower = file.name.lower()
            is_docx = file_name_lower.endswith('.docx') or file_name_lower.endswith('.doc')
            
            # Save uploaded file temporarily with chunked writing for memory efficiency
            file_ext = os.path.splitext(file.name)[1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                # Use chunked reading to avoid loading entire file into memory
                chunk_size = 8192  # 8KB chunks
                for chunk in file.chunks(chunk_size):
                    tmp_file.write(chunk)
                tmp_file_path = tmp_file.name

            try:
                # For DOCX files, convert to text first
                if is_docx:
                    text_content = self._extract_docx_text(tmp_file_path)
                    # Create a temporary text file for Gemini
                    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as text_file:
                        text_file.write(text_content)
                        text_file_path = text_file.name
                    
                    # Clean up original DOCX file
                    try:
                        os.unlink(tmp_file_path)
                    except Exception:
                        pass
                    tmp_file_path = text_file_path
                    file_ext = '.txt'

                # Upload file to Gemini for analysis (API uses 'file' parameter, not 'path')
                uploaded_file = self.client.files.upload(file=tmp_file_path)
                
                # Wait for file to be processed with exponential backoff and timeout
                uploaded_file = self._wait_for_file_processing(uploaded_file)

                # Check file state - handle both enum and string formats
                state = uploaded_file.state
                state_name = state.name if hasattr(state, 'name') else str(state)
                
                if state_name != "ACTIVE":
                    raise Exception(f"File processing failed: {state_name}")

                # Create prompt for metadata extraction
                prompt = """Analyze this educational content file and extract the following metadata in JSON format:
{
  "title": "A concise, descriptive title for this content (max 100 characters)",
  "description": "A brief description summarizing the content (2-3 sentences)",
  "subject": "The main subject or topic (e.g., Mathematics, Science, History, Literature, etc.)",
  "difficulty": "One of: beginner, intermediate, or advanced",
  "author": "Author name if mentioned, otherwise empty string",
  "publication_year": "Year of publication if mentioned, otherwise null"
}

Guidelines:
- Title: Extract from document title, header, or create a descriptive title based on content
- Description: Summarize the main topics and purpose in 2-3 sentences
- Subject: Identify the primary academic subject (be specific: Mathematics, Physics, Chemistry, Biology, History, Literature, etc.)
- Difficulty: Assess based on complexity, terminology, and concepts (beginner = introductory/elementary, intermediate = high school/undergraduate, advanced = graduate/professional)
- Author: Extract author name from document metadata or content if available
- Publication Year: Extract year from document if available

Return ONLY valid JSON, no additional text."""

                # Generate content with file using the file URI (API requires keyword arguments)
                file_part = types.Part.from_uri(
                    file_uri=uploaded_file.uri,
                    mime_type=uploaded_file.mime_type
                )
                
                # Create content parts: file + text prompt (API requires keyword argument)
                contents = [file_part, types.Part.from_text(text=prompt)]
                
                response = self.client.models.generate_content(
                    model=self.model,
                    contents=contents
                )

                # Parse JSON response with optimized string cleaning
                response_text = self._extract_json_from_response(response.text)
                metadata = json.loads(response_text)

                # Validate and normalize metadata (optimized with single-pass processing)
                result = self._normalize_metadata(metadata)

                return result

            finally:
                # Cleanup: Delete uploaded file from Gemini and local temp file
                self._cleanup_resources(uploaded_file, tmp_file_path)

        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse AI response as JSON: {str(e)}")
        except Exception as e:
            # Ensure cleanup even on error
            self._cleanup_resources(uploaded_file, tmp_file_path)
            raise Exception(f"Failed to extract metadata: {str(e)}")

    def _wait_for_file_processing(self, uploaded_file, max_attempts=None, timeout=None):
        """
        Wait for file processing with exponential backoff and timeout.
        Optimized to reduce unnecessary API calls.
        """
        max_attempts = max_attempts or self.MAX_POLLING_ATTEMPTS
        timeout = timeout or (max_attempts * self.MAX_POLL_INTERVAL)
        
        start_time = time.time()
        poll_interval = self.INITIAL_POLL_INTERVAL
        attempts = 0

        # Helper to get state name (handles both enum and string)
        def get_state_name(file_obj):
            state = file_obj.state
            if hasattr(state, 'name'):
                return state.name
            return str(state)

        # Check initial state
        current_state = get_state_name(uploaded_file)
        
        while current_state == "PROCESSING":
            # Check timeout
            if time.time() - start_time > timeout:
                raise Exception(f"File processing timeout after {timeout}s")
            
            # Check max attempts
            if attempts >= max_attempts:
                raise Exception(f"File processing exceeded {max_attempts} attempts")
            
            # Exponential backoff
            time.sleep(poll_interval)
            poll_interval = min(
                poll_interval * self.POLL_BACKOFF_MULTIPLIER,
                self.MAX_POLL_INTERVAL
            )
            
            # Refresh file status
            uploaded_file = self.client.files.get(uploaded_file.name)
            current_state = get_state_name(uploaded_file)
            attempts += 1

        return uploaded_file

    def _extract_json_from_response(self, response_text: str) -> str:
        """
        Extract JSON from response text efficiently.
        Handles markdown code blocks and plain JSON.
        """
        response_text = response_text.strip()
        
        # Try regex pattern matching first (most common case)
        match = self.JSON_PATTERN.search(response_text)
        if match:
            return match.group(1).strip()
        
        # Fallback: simple prefix/suffix removal
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        elif response_text.startswith('```'):
            response_text = response_text[3:]
        
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        
        return response_text.strip()

    def _normalize_metadata(self, metadata: dict) -> dict:
        """
        Normalize and validate metadata in a single pass.
        Uses efficient string operations and validation.
        """
        # Pre-define valid difficulties as set for O(1) lookup
        VALID_DIFFICULTIES = {'beginner', 'intermediate', 'advanced'}
        MIN_YEAR = 1900
        MAX_YEAR = 2100
        
        # Extract and normalize in one pass
        title = (metadata.get('title') or '').strip()[:500]
        description = (metadata.get('description') or '').strip()[:2000]
        subject = (metadata.get('subject') or '').strip()[:100]
        difficulty = (metadata.get('difficulty') or 'beginner').lower()
        author = (metadata.get('author') or '').strip()[:200]
        publication_year = metadata.get('publication_year')
        
        # Validate difficulty with set lookup (O(1))
        if difficulty not in VALID_DIFFICULTIES:
            difficulty = 'beginner'
        
        # Validate publication_year
        if publication_year:
            try:
                year = int(publication_year)
                if MIN_YEAR <= year <= MAX_YEAR:
                    publication_year = str(year)
                else:
                    publication_year = None
            except (ValueError, TypeError):
                publication_year = None
        
        return {
            'title': title,
            'description': description,
            'subject': subject,
            'difficulty': difficulty,
            'author': author,
            'publication_year': publication_year,
        }

    def _extract_docx_text(self, docx_path: str) -> str:
        """
        Extract text content from DOCX file.
        Returns plain text for Gemini processing.
        """
        try:
            from docx import Document
            
            doc = Document(docx_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
        except ImportError:
            raise Exception("python-docx library is required for DOCX files. Please install it: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    def _cleanup_resources(self, uploaded_file, tmp_file_path):
        """Clean up uploaded file and temporary file resources"""
        # Delete from Gemini
        if uploaded_file:
            try:
                self.client.files.delete(uploaded_file.name)
            except Exception:
                pass  # Ignore deletion errors
        
        # Delete local temp file
        if tmp_file_path and os.path.exists(tmp_file_path):
            try:
                os.unlink(tmp_file_path)
            except Exception:
                pass  # Ignore deletion errors

